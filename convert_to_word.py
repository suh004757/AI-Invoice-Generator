"""
Convert markdown files to Word documents
"""
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re
import os

def markdown_to_word(md_file, output_file):
    """Convert markdown file to Word document"""
    doc = Document()
    
    # Read markdown file
    with open(md_file, 'r', encoding='utf-8') as f:
        content = f.read()
    
    lines = content.split('\n')
    in_code_block = False
    code_language = ''
    code_lines = []
    
    for line in lines:
        # Code block detection
        if line.startswith('```'):
            if not in_code_block:
                # Start code block
                in_code_block = True
                code_language = line[3:].strip()
                code_lines = []
            else:
                # End code block
                in_code_block = False
                # Add code block to document
                if code_lines:
                    p = doc.add_paragraph()
                    p.style = 'Normal'
                    run = p.add_run('\n'.join(code_lines))
                    run.font.name = 'Consolas'
                    run.font.size = Pt(9)
                    run.font.color.rgb = RGBColor(0, 0, 0)
                    p.paragraph_format.left_indent = Inches(0.5)
                    # Light gray background (simulated with border)
                    p.paragraph_format.space_before = Pt(6)
                    p.paragraph_format.space_after = Pt(6)
                code_lines = []
            continue
        
        if in_code_block:
            code_lines.append(line)
            continue
        
        # Skip mermaid diagrams (not supported in Word)
        if 'mermaid' in line.lower() or line.strip().startswith('graph ') or line.strip().startswith('sequenceDiagram'):
            continue
        
        # Headers
        if line.startswith('# '):
            p = doc.add_heading(line[2:], level=1)
        elif line.startswith('## '):
            p = doc.add_heading(line[3:], level=2)
        elif line.startswith('### '):
            p = doc.add_heading(line[4:], level=3)
        elif line.startswith('#### '):
            p = doc.add_heading(line[5:], level=4)
        
        # Horizontal rule
        elif line.strip() == '---':
            p = doc.add_paragraph()
            p.add_run('_' * 80)
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
        
        # Lists
        elif line.strip().startswith('- ') or line.strip().startswith('* '):
            text = line.strip()[2:]
            # Remove markdown checkboxes
            text = re.sub(r'\[[ x/]\]\s*', '', text)
            p = doc.add_paragraph(text, style='List Bullet')
        
        elif re.match(r'^\d+\.\s', line.strip()):
            text = re.sub(r'^\d+\.\s', '', line.strip())
            p = doc.add_paragraph(text, style='List Number')
        
        # Blockquotes (alerts)
        elif line.strip().startswith('>'):
            text = line.strip()[1:].strip()
            # Remove alert markers
            text = re.sub(r'\[!(NOTE|TIP|IMPORTANT|WARNING|CAUTION)\]', '', text)
            if text:
                p = doc.add_paragraph(text)
                p.style = 'Quote'
        
        # Tables (basic support)
        elif '|' in line and line.strip().startswith('|'):
            # Skip table separator lines
            if re.match(r'\|[\s\-:]+\|', line):
                continue
            # Parse table row
            cells = [cell.strip() for cell in line.split('|')[1:-1]]
            if cells:
                # Create table if needed (simplified - just add as paragraph)
                p = doc.add_paragraph(' | '.join(cells))
                p.style = 'Normal'
        
        # Normal paragraphs
        elif line.strip():
            # Skip file links and image references
            if line.strip().startswith('![') or 'file:///' in line:
                continue
            
            # Remove markdown formatting
            text = line
            # Bold
            text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
            # Italic
            text = re.sub(r'\*(.+?)\*', r'\1', text)
            # Inline code
            text = re.sub(r'`(.+?)`', r'\1', text)
            # Links
            text = re.sub(r'\[(.+?)\]\(.+?\)', r'\1', text)
            
            if text.strip():
                p = doc.add_paragraph(text)
        
        # Empty line
        else:
            if doc.paragraphs and doc.paragraphs[-1].text.strip():
                doc.add_paragraph()
    
    # Save document
    doc.save(output_file)
    print(f"[OK] Created: {output_file}")

if __name__ == '__main__':
    # Artifact directory
    artifact_dir = r'C:\Users\tjwns\.gemini\antigravity\brain\3e96f4a6-f10f-4d61-b39a-d9810395495b'
    output_dir = r'c:\Users\tjwns\Desktop\NEW PROJECT\docs'
    
    # Create output directory
    os.makedirs(output_dir, exist_ok=True)
    
    # Files to convert
    files = [
        ('implementation_plan.md', 'AI_Invoice_Builder_Implementation_Plan.docx'),
        ('AI_ARCHITECTURE.md', 'AI_Invoice_Builder_Architecture.docx'),
        ('task.md', 'AI_Invoice_Builder_Tasks.docx')
    ]
    
    print("Converting markdown files to Word documents...\n")
    
    for md_file, docx_file in files:
        md_path = os.path.join(artifact_dir, md_file)
        docx_path = os.path.join(output_dir, docx_file)
        
        if os.path.exists(md_path):
            try:
                markdown_to_word(md_path, docx_path)
            except Exception as e:
                print(f"[ERROR] Error converting {md_file}: {e}")
        else:
            print(f"[ERROR] File not found: {md_path}")
    
    print(f"\n완료! Word 문서들이 '{output_dir}' 폴더에 저장되었습니다.")
    print(f"폴더 열기: {output_dir}")
