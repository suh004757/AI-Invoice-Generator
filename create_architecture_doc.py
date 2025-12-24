"""
Create improved Word document with embedded images
"""
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

def create_architecture_doc_with_images():
    """Create architecture document with embedded images"""
    doc = Document()
    
    # Title
    title = doc.add_heading('AI Invoice Builder - Architecture', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Overview
    doc.add_heading('System Overview', 1)
    p = doc.add_paragraph()
    p.add_run('Concept: ').bold = True
    p.add_run('"PO 던지면 인보이스 나오는 상자" (Throw in a PO, get an invoice out)')
    
    p = doc.add_paragraph()
    p.add_run('User Flow: ').bold = True
    p.add_run('Upload PO → Select Type → AI Analyze → Review → Generate Invoice')
    
    doc.add_paragraph()
    
    # System Architecture Diagram
    doc.add_heading('System Architecture Diagram', 1)
    doc.add_paragraph('전체 시스템의 계층 구조와 컴포넌트 간 연결을 보여줍니다.')
    
    img_path = r'C:\Users\tjwns\.gemini\antigravity\brain\3e96f4a6-f10f-4d61-b39a-d9810395495b\system_architecture_diagram_1765451659641.png'
    if os.path.exists(img_path):
        doc.add_picture(img_path, width=Inches(6.5))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_page_break()
    
    # User Flow Sequence
    doc.add_heading('User Flow Sequence', 1)
    doc.add_paragraph('사용자가 PO를 업로드하고 인보이스를 생성하는 전체 흐름입니다.')
    
    img_path = r'C:\Users\tjwns\.gemini\antigravity\brain\3e96f4a6-f10f-4d61-b39a-d9810395495b\user_flow_sequence_1765451681298.png'
    if os.path.exists(img_path):
        doc.add_picture(img_path, width=Inches(6.5))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_page_break()
    
    # Database Schema
    doc.add_heading('Database Schema (ERD)', 1)
    doc.add_paragraph('데이터베이스 테이블 구조와 관계를 보여줍니다.')
    
    img_path = r'C:\Users\tjwns\.gemini\antigravity\brain\3e96f4a6-f10f-4d61-b39a-d9810395495b\database_schema_erd_1765451704233.png'
    if os.path.exists(img_path):
        doc.add_picture(img_path, width=Inches(6.5))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_page_break()
    
    # Technology Stack
    doc.add_heading('Technology Stack', 1)
    
    table = doc.add_table(rows=11, cols=3)
    table.style = 'Light Grid Accent 1'
    
    # Header
    header_cells = table.rows[0].cells
    header_cells[0].text = 'Layer'
    header_cells[1].text = 'Technology'
    header_cells[2].text = 'Purpose'
    
    # Data
    data = [
        ('GUI Framework', 'PySide6 (Qt6)', 'Native Windows interface'),
        ('AI/LLM', 'Anthropic Claude API', 'Primary AI extraction'),
        ('AI/LLM', 'OpenAI API', 'Alternative AI provider'),
        ('AI/LLM', 'LM Studio / Ollama', 'Local LLM support'),
        ('PDF Processing', 'pdfplumber', 'PDF text extraction'),
        ('OCR', 'pytesseract / EasyOCR', 'Image text extraction'),
        ('Image Processing', 'Pillow', 'Image manipulation'),
        ('Database', 'SQLite3', 'Data storage'),
        ('Excel Processing', 'openpyxl', 'Excel generation'),
        ('Language', 'Python 3.11+', 'Application logic'),
    ]
    
    for i, (layer, tech, purpose) in enumerate(data, start=1):
        cells = table.rows[i].cells
        cells[0].text = layer
        cells[1].text = tech
        cells[2].text = purpose
    
    doc.add_paragraph()
    
    # Key Features
    doc.add_heading('Key Features', 1)
    
    features = [
        ('Multi-Provider LLM Support', 'Claude API, OpenAI API, Local LLM (LM Studio/Ollama)'),
        ('Intelligent Document Processing', 'PDF text extraction, OCR for scanned documents, Image processing'),
        ('Confidence Scoring', 'Per-field confidence indicators, Visual highlighting of low-confidence fields'),
        ('Flexible Review', 'Side-by-side PO viewer, Editable fields, Quick corrections'),
        ('JSON Metadata Storage', 'Flexible schema for varying PO fields, AI-extracted data preservation'),
    ]
    
    for title, desc in features:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(title + ': ').bold = True
        p.add_run(desc)
    
    doc.add_paragraph()
    
    # File Organization
    doc.add_heading('File Organization', 1)
    
    p = doc.add_paragraph('invoice_studio/', style='Normal')
    p.paragraph_format.left_indent = Inches(0)
    
    structure = [
        ('ai/', 'NEW: AI/LLM integration'),
        ('  llm_client.py', 'LLM client factory & implementations'),
        ('  prompts.py', 'Prompt templates'),
        ('  extractor.py', 'Main extraction orchestrator'),
        ('document/', 'NEW: Document processing'),
        ('  pdf_processor.py', 'PDF handling'),
        ('  image_processor.py', 'Image & OCR'),
        ('  text_processor.py', 'Text cleaning'),
        ('gui/wizard/', 'NEW: Wizard UI'),
        ('  upload_step.py', 'Step 1: Upload PO'),
        ('  review_step.py', 'Step 2: Review extracted data'),
        ('  generate_step.py', 'Step 3: Generate invoice'),
        ('models/', 'MODIFIED: Business models'),
        ('  invoice.py', 'Add from_extracted_data()'),
        ('  purchase_order.py', 'NEW: PO model'),
        ('database/', 'MODIFIED: Database layer'),
        ('  models.py', 'Add PO and extraction_logs tables'),
        ('config.py', 'NEW: Configuration management'),
    ]
    
    for item, desc in structure:
        p = doc.add_paragraph(style='Normal')
        run1 = p.add_run('├── ' + item)
        run1.font.name = 'Consolas'
        run1.font.size = Pt(9)
        p.add_run('  # ' + desc)
        p.paragraph_format.left_indent = Inches(0.3)
    
    doc.add_paragraph()
    
    # Configuration
    doc.add_heading('Configuration (.env)', 1)
    
    p = doc.add_paragraph()
    code_text = """# LLM Provider Selection
LLM_PROVIDER=claude  # Options: claude, openai, lm_studio

# API Keys
CLAUDE_API_KEY=sk-ant-xxxxx
OPENAI_API_KEY=sk-xxxxx

# LM Studio Configuration
LM_STUDIO_URL=http://localhost:1234
LM_STUDIO_MODEL=local-model

# Application Defaults
DEFAULT_CURRENCY=KRW
DEFAULT_INVOICE_TYPE=tax"""
    
    run = p.add_run(code_text)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    p.paragraph_format.left_indent = Inches(0.5)
    
    # Save
    output_path = r'c:\Users\tjwns\Desktop\NEW PROJECT\docs\AI_Invoice_Builder_Architecture_with_Diagrams.docx'
    doc.save(output_path)
    print(f"[OK] Created: {output_path}")
    return output_path

if __name__ == '__main__':
    output_path = create_architecture_doc_with_images()
    print(f"\n완료! 다이어그램이 포함된 문서가 생성되었습니다.")
    print(f"파일: {output_path}")
